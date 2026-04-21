#!/usr/bin/env python3
"""Convert proposal-he.md → proposal-he.docx with embedded images and RTL."""

import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
MD_FILE   = os.path.join(BASE_DIR, 'proposal-he.md')
OUT_FILE  = os.path.join(BASE_DIR, 'proposal-he.docx')
IMG_DIR   = os.path.join(BASE_DIR, '..', 'screenshots')

# ── helpers ──────────────────────────────────────────────────────────────────

def rtl(para, align=WD_ALIGN_PARAGRAPH.RIGHT):
    pPr = para._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    para.paragraph_format.alignment = align

def shade_para(para, fill='F2F2F2'):
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def hr(doc):
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bot)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)

def cell_shade(cell, fill):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

# Inline markdown: **bold**, *italic*, `code`, plain text
_INLINE = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+)', re.DOTALL)

def add_runs(para, text, base_size=None):
    for m in _INLINE.finditer(text):
        bold, italic, code, plain = m.group(1), m.group(2), m.group(3), m.group(4)
        if bold:
            r = para.add_run(bold); r.bold = True
            if base_size: r.font.size = Pt(base_size)
        elif italic:
            r = para.add_run(italic); r.italic = True
            if base_size: r.font.size = Pt(base_size)
        elif code:
            r = para.add_run(code)
            r.font.name = 'Courier New'; r.font.size = Pt(9)
        elif plain:
            r = para.add_run(plain)
            if base_size: r.font.size = Pt(base_size)

def resolve_img(md_path):
    # md_path is like ../screenshots/01_home.png
    raw = md_path.strip().lstrip('./')
    # try direct relative resolution from docs/
    candidate = os.path.normpath(os.path.join(BASE_DIR, md_path))
    if os.path.exists(candidate):
        return candidate
    # try just the filename in IMG_DIR
    fname = os.path.basename(md_path)
    candidate2 = os.path.join(IMG_DIR, fname)
    if os.path.exists(candidate2):
        return candidate2
    return None

# ── flush helpers ─────────────────────────────────────────────────────────────

def flush_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'
    for ri, row_data in enumerate(rows):
        for ci in range(cols):
            cell_text = row_data[ci] if ci < len(row_data) else ''
            cell = tbl.cell(ri, ci)
            p = cell.paragraphs[0]
            add_runs(p, cell_text, base_size=10)
            rtl(p)
            if ri == 0:                         # header row
                for run in p.runs:
                    run.bold = True
                cell_shade(cell, 'D9E2F3')      # blue tint header
            elif ri % 2 == 0:
                cell_shade(cell, 'F7F8FC')      # alternating row
    doc.add_paragraph()                          # breathing room after table

def flush_code(doc, lines):
    if not lines:
        return
    text = '\n'.join(lines)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    shade_para(para, 'F4F4F4')
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.5)

def flush_list(doc, items, numbered):
    style = 'List Number' if numbered else 'List Bullet'
    for (text, level) in items:
        para = doc.add_paragraph(style=style)
        add_runs(para, text)
        rtl(para)
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)

# ── main converter ────────────────────────────────────────────────────────────

def convert():
    doc = Document()

    # Page setup: A4, sensible margins
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin   = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Default font supporting Hebrew
    for sname in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                  'List Bullet', 'List Number'):
        try:
            s = doc.styles[sname]
            s.font.name = 'Arial'
        except Exception:
            pass

    # Heading colours
    for lvl, colour in [(1, '1F3864'), (2, '2E5496'), (3, '2E74B5')]:
        try:
            s = doc.styles[f'Heading {lvl}']
            s.font.color.rgb = RGBColor(*bytes.fromhex(colour))
            s.font.size = Pt([20, 15, 12][lvl - 1])
        except Exception:
            pass

    # ── parse ────────────────────────────────────────────────────────────────
    with open(MD_FILE, encoding='utf-8') as f:
        lines = f.read().splitlines()

    i = 0
    in_code  = False
    code_buf = []
    in_table = False
    tbl_rows = []
    list_buf = []          # [(text, level)]
    list_num = False

    def flush_pending_list():
        nonlocal list_buf, list_num
        if list_buf:
            flush_list(doc, list_buf, list_num)
            list_buf = []

    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()
        i  += 1

        # ── HTML wrappers ────────────────────────────────────────────────────
        if s.startswith('<div') or s == '</div>':
            continue

        # ── code fence ───────────────────────────────────────────────────────
        if s.startswith('```'):
            flush_pending_list()
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                in_code = False
                flush_code(doc, code_buf)
                code_buf = []
            continue

        if in_code:
            code_buf.append(raw)
            continue

        # ── table row ────────────────────────────────────────────────────────
        if s.startswith('|'):
            flush_pending_list()
            if not in_table:
                in_table = True
                tbl_rows = []
            if re.match(r'^\|[\s\-\|:]+\|$', s):   # separator row
                continue
            cells = [c.strip() for c in s.split('|')]
            cells = [c for c in cells if c != '']
            if cells:
                tbl_rows.append(cells)
            continue
        else:
            if in_table:
                in_table = False
                flush_pending_list()
                flush_table(doc, tbl_rows)
                tbl_rows = []

        # ── blank line ───────────────────────────────────────────────────────
        if s == '':
            flush_pending_list()
            continue

        # ── horizontal rule ──────────────────────────────────────────────────
        if re.match(r'^-{3,}$', s) or re.match(r'^={3,}$', s):
            flush_pending_list()
            hr(doc)
            continue

        # ── headings ─────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,3}) (.+)$', s)
        if m:
            flush_pending_list()
            lvl  = len(m.group(1))
            text = m.group(2)
            para = doc.add_heading(text, level=lvl)
            rtl(para)
            continue

        # ── image ────────────────────────────────────────────────────────────
        m = re.match(r'!\[.*?\]\((.+?)\)', s)
        if m:
            flush_pending_list()
            path = resolve_img(m.group(1))
            if path:
                try:
                    para = doc.add_paragraph()
                    rtl(para, WD_ALIGN_PARAGRAPH.CENTER)
                    para.add_run().add_picture(path, width=Cm(14))
                    # caption
                    alt_m = re.match(r'!\[(.+?)\]', s)
                    if alt_m and alt_m.group(1):
                        cap = doc.add_paragraph(alt_m.group(1))
                        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in cap.runs:
                            r.font.size = Pt(9)
                            r.italic = True
                            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                except Exception as e:
                    p = doc.add_paragraph(f'[תמונה: {os.path.basename(path)}]')
                    rtl(p)
            continue

        # ── unordered list ───────────────────────────────────────────────────
        m = re.match(r'^[-*] (.+)$', s)
        if m:
            if list_buf and list_num:
                flush_pending_list()
            list_num = False
            list_buf.append((m.group(1), 0))
            continue

        # ── ordered list ─────────────────────────────────────────────────────
        m = re.match(r'^\d+\. (.+)$', s)
        if m:
            if list_buf and not list_num:
                flush_pending_list()
            list_num = True
            list_buf.append((m.group(1), 0))
            continue

        # ── blockquote ───────────────────────────────────────────────────────
        if s.startswith('>'):
            flush_pending_list()
            text = re.sub(r'^>+\s*', '', s)
            para = doc.add_paragraph()
            add_runs(para, text, base_size=11)
            rtl(para)
            para.paragraph_format.left_indent  = Cm(0.8)
            para.paragraph_format.right_indent = Cm(0.0)
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                run.italic = True
            shade_para(para, 'EEF2F8')
            continue

        # ── regular paragraph ────────────────────────────────────────────────
        flush_pending_list()
        para = doc.add_paragraph()
        add_runs(para, s)
        rtl(para)

    # flush anything remaining
    flush_pending_list()
    if in_table and tbl_rows:
        flush_table(doc, tbl_rows)
    if in_code and code_buf:
        flush_code(doc, code_buf)

    doc.save(OUT_FILE)
    print(f'✓  Saved → {OUT_FILE}')

if __name__ == '__main__':
    convert()
